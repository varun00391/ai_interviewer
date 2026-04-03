import io
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


def extract_text_from_file(path: str) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(path)
        parts: list[str] = []
        for page in reader.pages:
            t = page.extract_text() or ""
            parts.append(t)
        return "\n".join(parts).strip()
    if suffix in (".docx", ".doc"):
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs).strip()
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        parts = [(page.extract_text() or "") for page in reader.pages]
        return "\n".join(parts).strip()
    if name.endswith(".docx"):
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    return data.decode("utf-8", errors="ignore")
