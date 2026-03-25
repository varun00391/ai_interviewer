import io
import logging
import zipfile
import xml.etree.ElementTree as ET
from typing import BinaryIO

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

logger = logging.getLogger(__name__)

_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _is_ole_compound_file(data: bytes) -> bool:
    return len(data) >= len(_OLE_MAGIC) and data[: len(_OLE_MAGIC)] == _OLE_MAGIC


def _is_word_docx_package(data: bytes) -> bool:
    """True if bytes are a ZIP OOXML package containing a Word main document part."""
    if len(data) < 4 or data[:2] != b"PK":
        return False
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            return "word/document.xml" in z.namelist()
    except zipfile.BadZipFile:
        return False


def extract_resume_text(*, filename: str, fileobj: BinaryIO) -> str:
    raw = fileobj.read()
    if not raw:
        raise ValueError("Empty file")
    ext = (filename or "").rsplit(".", 1)[-1].lower()

    if len(raw) >= 4 and raw[:4] == b"%PDF":
        return _pdf_text(raw)

    if ext == "docx":
        if not _is_word_docx_package(raw):
            raise ValueError(
                "This file is not a valid Word .docx (Open XML). Export again from Word or Google Docs as .docx."
            )
        return _docx_text(raw)

    sniffed = _sniff_format(raw)
    if sniffed == "docx":
        if _is_word_docx_package(raw):
            return _docx_text(raw)
        # ZIP but not Word — do not mis-handle as docx

    if ext == "doc" or (_is_ole_compound_file(raw) and sniffed != "docx"):
        raise ValueError("Legacy .doc is not supported. Save as .docx in Word and upload again.")

    if sniffed == "pdf" or ext == "pdf":
        return _pdf_text(raw)

    raise ValueError("Unsupported file type. Upload a .pdf or .docx file.")


def _sniff_format(data: bytes) -> str | None:
    """Infer PDF vs ZIP when the browser sends a generic MIME type or odd filename."""
    if len(data) >= 4 and data[:4] == b"%PDF":
        return "pdf"
    if len(data) >= 2 and data[:2] == b"PK":
        try:
            if zipfile.is_zipfile(io.BytesIO(data)):
                return "docx"
        except Exception:
            pass
    return None


def _pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data), strict=False)
    except PdfReadError as e:
        logger.warning("PdfReadError: %s", e)
        raise ValueError("Could not read this PDF (it may be corrupted or password-protected).") from e
    parts: list[str] = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            parts.append(t)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Could not extract text from PDF (it may be image-only).")
    return text


def _paragraphs_and_tables_from_block(container) -> list[str]:
    """Collect text from a python-docx Document, Header, Footer, or Cell."""
    parts: list[str] = []
    try:
        for p in container.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
    except Exception:
        pass
    try:
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.extend(_paragraphs_and_tables_from_block(cell))
    except Exception:
        pass
    return parts


def _docx_text_python_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    parts.extend(_paragraphs_and_tables_from_block(doc))
    for section in doc.sections:
        for attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            try:
                block = getattr(section, attr, None)
                if block is not None:
                    parts.extend(_paragraphs_and_tables_from_block(block))
            except Exception:
                continue
    return "\n".join(parts).strip()


def _docx_xml_fallback_text(data: bytes) -> str:
    """
    Pull all w:t text from document, headers, and footers. Helps when content lives in
    structures python-docx does not fully expose (some templates, text boxes in body XML, etc.).
    """
    chunks: list[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            names = z.namelist()
            targets = []
            for n in names:
                if n == "word/document.xml":
                    targets.append(n)
                elif n.startswith("word/header") and n.endswith(".xml"):
                    targets.append(n)
                elif n.startswith("word/footer") and n.endswith(".xml"):
                    targets.append(n)
            for name in sorted(set(targets)):
                try:
                    xml_bytes = z.read(name)
                    root = ET.fromstring(xml_bytes)
                except (KeyError, ET.ParseError) as e:
                    logger.debug("Skip part %s: %s", name, e)
                    continue
                for el in root.iter():
                    tag = el.tag
                    if tag == f"{_W_NS}t":
                        if el.text:
                            chunks.append(el.text)
                    elif tag in (f"{_W_NS}tab",):
                        chunks.append("\t")
                    elif tag in (f"{_W_NS}br", f"{_W_NS}cr"):
                        chunks.append("\n")
    except Exception as e:
        logger.warning("docx XML fallback failed: %s", e)
        return ""
    return " ".join(chunks).replace("\t", " ").strip()


def _docx_text(data: bytes) -> str:
    try:
        primary = _docx_text_python_docx(data)
    except Exception as e:
        logger.warning("python-docx could not open document: %s", e)
        primary = ""

    if len(primary) < 40:
        fallback = _docx_xml_fallback_text(data)
        merged = "\n".join(x for x in (primary, fallback) if x).strip()
        if merged:
            return merged

    if primary:
        return primary

    fallback = _docx_xml_fallback_text(data)
    if fallback.strip():
        return fallback.strip()

    raise ValueError(
        "Could not extract text from this .docx. Try copying the resume into a new blank Word document "
        "and saving again, or export as PDF."
    )
